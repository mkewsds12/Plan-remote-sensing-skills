"""
空天地一体化监测技术方案 - Markdown 转 Word 文档工具

用法：
    python md2docx.py <输入markdown文件> [输出docx文件]

示例：
    python md2docx.py 林业生态监测方案.md
    python md2docx.py 方案.md 输出方案.docx

依赖：
    pip install python-docx
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn


# ── 样式配置 ──────────────────────────────────────────────
# 字体配置（按党政机关公文格式规范 GB/T 9704-2012）
# - 标题：黑体加粗
# - 正文/表格：仿宋，西文用 Times New Roman
# 字号对照：
#   Pt(22) = 二号（封面标题）
#   Pt(18) = 小二号（一级标题，如"1. 总体规划与建设目标"）
#   Pt(16) = 三号（二级标题，如"1.1 项目背景"）
#   Pt(15) = 小三号（三级标题，如"1.1.1 政策背景"）
#   Pt(14) = 四号（四级标题 + 正文）
#   Pt(12) = 小四号（表格内文字）
#   Pt(28) = 行距 28 磅固定值（约 0.99cm，公文标准行距）

FONT_CONFIG = {
    "title_cn": "方正小标宋简体",   # 封面标题字体（预留）
    "heading_cn": "黑体",             # 各级标题中文字体
    "body_cn": "仿宋",               # 正文中文字体
    "body_en": "Times New Roman",    # 正文/标题西文字体
    "font_size_title": Pt(22),       # 封面标题：二号
    "font_size_h1": Pt(18),          # 一级标题：小二号
    "font_size_h2": Pt(16),          # 二级标题：三号
    "font_size_h3": Pt(15),          # 三级标题：小三号
    "font_size_h4": Pt(14),          # 四级标题：四号
    "font_size_body": Pt(14),        # 正文：四号（仿宋 14 磅）
    "font_size_table": Pt(12),       # 表格内文字：小四号
    "line_spacing": Pt(28),          # 行距：28 磅固定行距
}

# 页面边距配置（按公文标准 GB/T 9704-2012）
# 上 3.7cm、下 3.5cm、左 2.8cm（装订侧稍宽）、右 2.6cm
PAGE_CONFIG = {
    "top_margin": Cm(3.7),           # 上边距
    "bottom_margin": Cm(3.5),        # 下边距
    "left_margin": Cm(2.8),          # 左边距（装订侧）
    "right_margin": Cm(2.6),         # 右边距
}


def create_document():
    """创建并配置Word文档基础样式"""
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.top_margin = PAGE_CONFIG["top_margin"]
    section.bottom_margin = PAGE_CONFIG["bottom_margin"]
    section.left_margin = PAGE_CONFIG["left_margin"]
    section.right_margin = PAGE_CONFIG["right_margin"]

    # 默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_CONFIG["body_en"]
    font.size = FONT_CONFIG["font_size_body"]
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CONFIG["body_cn"])
    pf = style.paragraph_format
    pf.line_spacing = FONT_CONFIG["line_spacing"]
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(0.74)  # 两个字符缩进

    return doc


def set_run_font(run, cn_font=None, en_font=None, size=None, bold=False, color=None):
    """统一设置 run 的字体属性"""
    if en_font:
        run.font.name = en_font
    if cn_font:
        run.element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)
        # 让中文标点（如双引号""）也使用中文字体，而非 Times New Roman
        run.element.rPr.rFonts.set(qn("w:hAnsi"), cn_font)
    if size:
        run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading_styled(doc, text, level):
    """添加带样式的标题"""
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = FONT_CONFIG["line_spacing"]
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.first_line_indent = Cm(0)

    size_map = {
        1: FONT_CONFIG["font_size_h1"],
        2: FONT_CONFIG["font_size_h2"],
        3: FONT_CONFIG["font_size_h3"],
        4: FONT_CONFIG["font_size_h4"],
    }
    size = size_map.get(level, FONT_CONFIG["font_size_body"])

    if level == 1:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run(text)
    set_run_font(
        run,
        cn_font=FONT_CONFIG["heading_cn"],
        en_font=FONT_CONFIG["body_en"],
        size=size,
        bold=True,
    )

    # 设置大纲级别
    pPr = para._element.get_or_add_pPr()
    outline_lvl = pPr.makeelement(qn("w:outlineLvl"), {})
    outline_lvl.set(qn("w:val"), str(level - 1))
    pPr.append(outline_lvl)

    return para


def add_body_paragraph(doc, text, indent=True):
    """添加正文段落，解析行内加粗/斜体"""
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = FONT_CONFIG["line_spacing"]
    if not indent:
        para.paragraph_format.first_line_indent = Cm(0)

    _parse_inline(para, text)
    return para


def add_bullet_item(doc, text, level=0):
    """添加列表项"""
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = FONT_CONFIG["line_spacing"]
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.left_indent = Cm(0.74 * (level + 1))
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)

    # 添加项目符号
    bullets = ["● ", "○ ", "■ "]
    prefix = bullets[min(level, len(bullets) - 1)]

    _parse_inline(para, prefix + text)
    return para


def add_table(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Cm(0)
        run = para.add_run(header.strip())
        set_run_font(
            run,
            cn_font=FONT_CONFIG["heading_cn"],
            en_font=FONT_CONFIG["body_en"],
            size=FONT_CONFIG["font_size_table"],
            bold=True,
        )

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= len(headers):
                break
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Cm(0)
            run = para.add_run(cell_text.strip())
            set_run_font(
                run,
                cn_font=FONT_CONFIG["body_cn"],
                en_font=FONT_CONFIG["body_en"],
                size=FONT_CONFIG["font_size_table"],
            )

    # 表格后空一行
    doc.add_paragraph()
    return table


def add_image_placeholder(doc, text):
    """添加图表占位符"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run(text)
    set_run_font(
        run,
        cn_font=FONT_CONFIG["body_cn"],
        en_font=FONT_CONFIG["body_en"],
        size=FONT_CONFIG["font_size_body"],
        color=RGBColor(128, 128, 128),
    )
    return para


def _parse_inline(para, text):
    """解析行内 Markdown 格式（加粗、斜体、行内代码）"""
    # 正则匹配加粗、斜体、行内代码
    pattern = r"(\*\*(.+?)\*\*|__(.+?)__|`(.+?)`|\*(.+?)\*|_(.+?)_)"
    last_end = 0

    for match in re.finditer(pattern, text):
        # 添加匹配前的普通文本
        if match.start() > last_end:
            run = para.add_run(text[last_end : match.start()])
            set_run_font(
                run,
                cn_font=FONT_CONFIG["body_cn"],
                en_font=FONT_CONFIG["body_en"],
                size=FONT_CONFIG["font_size_body"],
            )

        # 加粗（仿宋加粗，不使用黑体）
        if match.group(2) or match.group(3):
            content = match.group(2) or match.group(3)
            run = para.add_run(content)
            set_run_font(
                run,
                cn_font=FONT_CONFIG["body_cn"],
                en_font=FONT_CONFIG["body_en"],
                size=FONT_CONFIG["font_size_body"],
                bold=True,
            )
        # 行内代码
        elif match.group(4):
            run = para.add_run(match.group(4))
            set_run_font(
                run,
                cn_font=FONT_CONFIG["body_cn"],
                en_font="Consolas",
                size=FONT_CONFIG["font_size_table"],
            )
        # 斜体
        elif match.group(5) or match.group(6):
            content = match.group(5) or match.group(6)
            run = para.add_run(content)
            set_run_font(
                run,
                cn_font=FONT_CONFIG["body_cn"],
                en_font=FONT_CONFIG["body_en"],
                size=FONT_CONFIG["font_size_body"],
            )
            run.font.italic = True

        last_end = match.end()

    # 剩余文本
    if last_end < len(text):
        run = para.add_run(text[last_end:])
        set_run_font(
            run,
            cn_font=FONT_CONFIG["body_cn"],
            en_font=FONT_CONFIG["body_en"],
            size=FONT_CONFIG["font_size_body"],
        )


# ── Markdown 解析 ────────────────────────────────────────

def parse_markdown(md_text):
    """将 Markdown 文本解析为结构化元素列表"""
    lines = md_text.split("\n")
    elements = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 空行
        if not stripped:
            i += 1
            continue

        # YAML front matter
        if stripped == "---" and i == 0:
            i += 1
            while i < len(lines) and lines[i].strip() != "---":
                i += 1
            i += 1
            continue

        # 引用块 (> ...) - 转为正文
        if stripped.startswith(">"):
            text = stripped.lstrip("> ").strip()
            if text:
                elements.append(("body", text, False))
            i += 1
            continue

        # 水平线
        if stripped in ("---", "***", "___"):
            i += 1
            continue

        # 标题
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            elements.append(("heading", text, level))
            i += 1
            continue

        # 图表占位符 [图：XXX] 或 [表：XXX]
        placeholder_match = re.match(r"^\[([图表]：.+?)\]$", stripped)
        if placeholder_match:
            elements.append(("placeholder", placeholder_match.group(1), None))
            i += 1
            continue

        # 表格（| ... | 开头）
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) >= 2:
                headers = [c.strip() for c in table_lines[0].split("|") if c.strip()]
                rows = []
                for tl in table_lines[2:]:  # 跳过分隔行
                    row = [c.strip() for c in tl.split("|") if c.strip()]
                    if row:
                        rows.append(row)
                elements.append(("table", headers, rows))
            continue

        # 列表项
        list_match = re.match(r"^(\s*)([-*+]|\d+[.)]) (.+)$", stripped)
        if list_match:
            indent = len(line) - len(line.lstrip())
            level = indent // 2
            text = list_match.group(3).strip()
            elements.append(("bullet", text, level))
            i += 1
            continue

        # 复选框
        checkbox_match = re.match(r"^- \[[ x]\] (.+)$", stripped)
        if checkbox_match:
            elements.append(("bullet", checkbox_match.group(1), 0))
            i += 1
            continue

        # 代码块 - 跳过
        if stripped.startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1
            continue

        # 普通段落
        elements.append(("body", stripped, True))
        i += 1

    return elements


def convert_md_to_docx(md_path, docx_path=None):
    """主转换函数"""
    md_path = Path(md_path)
    if docx_path is None:
        docx_path = md_path.with_suffix(".docx")
    else:
        docx_path = Path(docx_path)

    md_text = md_path.read_text(encoding="utf-8")
    elements = parse_markdown(md_text)

    doc = create_document()

    for elem_type, *args in elements:
        if elem_type == "heading":
            text, level = args
            add_heading_styled(doc, text, min(level, 4))
        elif elem_type == "body":
            text, indent = args
            add_body_paragraph(doc, text, indent=indent)
        elif elem_type == "bullet":
            text, level = args
            add_bullet_item(doc, text, level)
        elif elem_type == "table":
            headers, rows = args
            add_table(doc, headers, rows)
        elif elem_type == "placeholder":
            text = args[0]
            add_image_placeholder(doc, f"[{text}]")

    doc.save(str(docx_path))
    return docx_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    result = convert_md_to_docx(input_file, output_file)
    print(f"已生成: {result}")
